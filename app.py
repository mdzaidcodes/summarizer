from flask import Flask, render_template, request, send_file
from flask_socketio import SocketIO, emit
from sentence_transformers import SentenceTransformer
import pdfplumber
import os
import tempfile
import ollama
from docx import Document  # Import for reading .docx files
import traceback  # To help with detailed error logging

app = Flask(__name__)
socketio = SocketIO(app)

# Load the embedding model for creating embeddings
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Global placeholder for document text
document_text = ""
summary_text = ""

# Load the Llama model for response generation
model_name = 'llama3.1'
try:
    ollama.pull(model_name)
except Exception as e:
    print(f"Error preloading model '{model_name}': {e}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_document():
    # Check if the document file is in the request
    if 'document' not in request.files:
        return "No file uploaded", 400

    file = request.files['document']
    if file.filename == '':
        return "No file selected", 400

    # Save the uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        file.save(temp_file.name)
        temp_path = temp_file.name

    try:
        # Process the file based on its extension
        global document_text
        document_text = ""
        
        print(f"Attempting to extract text from: {file.filename}")
        
        if file.filename.lower().endswith('.pdf'):
            # Extract text from a PDF
            try:
                with pdfplumber.open(temp_path) as pdf:
                    document_text = " ".join(page.extract_text() for page in pdf.pages if page.extract_text())
                print("PDF text extraction successful.")
            except Exception as e:
                print("Error extracting text from PDF:", e)
                return "Error extracting text from PDF", 500

        elif file.filename.lower().endswith('.docx'):
            # Extract text from a Word document
            try:
                doc = Document(temp_path)
                document_text = " ".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
                print("DOCX text extraction successful.")
            except Exception as e:
                print("Error extracting text from DOCX:", e)
                return "Error extracting text from DOCX", 500
        else:
            return "Unsupported file format. Please upload a PDF or Word document.", 400

        if not document_text:
            print("No text extracted from the document.")
            return "Failed to extract text from the document. Please check the file content.", 400

        # Print the extracted text to the terminal for verification
        print("Extracted Text from Document:")
        print(document_text)
        print("-" * 50)  # Separator for readability

        return "Document uploaded and ready for summarization", 200
    except Exception as e:
        print("Error processing document:", e)
        traceback.print_exc()  # Print the traceback for detailed debugging
        return f"Error processing document: {str(e)}", 500
    finally:
        os.remove(temp_path)  # Clean up the temporary file

@app.route('/delete', methods=['POST'])
def delete_document():
    global document_text
    document_text = ""  # Clear the document text from memory
    return "Document deleted successfully", 200

@socketio.on('summarize')
def handle_summarization(data=None):  # Set a default value for data
    # Check if there is an uploaded document
    if not document_text:
        emit('summary', {'text': "No document uploaded or processed yet."})
        return

    # Updated prompt to clarify that the model is working with extracted text, not a full document
    prompt = (
        f"You are an AI assistant. The following text has been extracted from a document provided by the user. "
        f"The summar should not be more than 200 words."
        f"Please read this text carefully and generate a concise summary of its main points.\n\n"
        f"--- Extracted Text ---\n\n{document_text}\n\n"
        f"--- Summary ---"
    )

    try:
        response = ollama.chat(
            model=model_name,
            messages=[
                {'role': 'system', 'content': prompt},
                {'role': 'user', 'content': "Please summarize the extracted text provided above."}
            ],
        )

        # Extract the generated summary
        summary = response['message']['content']
        emit('summary', {'text': summary})
    except Exception as e:
        print("Error generating summary:", e)
        traceback.print_exc()
        emit('summary', {'text': f'Error generating summary: {str(e)}'})

@app.route('/download_summary', methods=['GET'])
def download_summary():
    global summary_text

    # Use the placeholder if no summary has been generated
    content = summary_text if summary_text else "Summary will appear here after uploading a document..."

    # Create a temporary .docx file
    temp_path = tempfile.mktemp(suffix=".docx")
    doc = Document()
    doc.add_paragraph(content)
    doc.save(temp_path)

    return send_file(temp_path, as_attachment=True, download_name="summary.docx")


if __name__ == '__main__':
    socketio.run(app, debug=True, use_reloader=False)
